"""Document parsing utilities for resume extraction.

Supports PDF, DOCX, and TXT file formats.
"""

import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path: str, mime_type: str | None = None) -> str:
    """Extract text content from a document.

    Args:
        file_path: Path to the file
        mime_type: MIME type hint (optional)

    Returns:
        Extracted text content
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext == ".txt" or mime_type == "text/plain":
        return extract_text_from_txt(file_path)
    elif ext == ".pdf" or mime_type == "application/pdf":
        return extract_text_from_pdf(file_path)
    elif ext in (".docx", ".doc") or "word" in (mime_type or ""):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def extract_text_from_txt(file_path: str) -> str:
    """Extract text from a plain text file."""
    try:
        with open(file_path, encoding="utf-8", errors="replace") as f:
            return f.read().strip()
    except Exception as e:
        logger.error(f"Failed to read TXT file {file_path}: {e}")
        raise ValueError(f"Failed to read text file: {e}") from e


def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from a PDF file using pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("pypdf not installed, PDF extraction unavailable")
        return "[PDF extraction requires pypdf package]"

    try:
        reader = PdfReader(file_path)
        text_parts = []

        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.error(f"Failed to extract PDF text from {file_path}: {e}")
        raise ValueError(f"Failed to extract PDF text: {e}") from e


def extract_text_from_docx(file_path: str) -> str:
    """Extract text from a DOCX file using python-docx."""
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx not installed, DOCX extraction unavailable")
        return "[DOCX extraction requires python-docx package]"

    try:
        doc = Document(file_path)
        text_parts = []

        # Extract from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))

        return "\n".join(text_parts).strip()
    except Exception as e:
        logger.error(f"Failed to extract DOCX text from {file_path}: {e}")
        raise ValueError(f"Failed to extract DOCX text: {e}") from e


def clean_extracted_text(text: str) -> str:
    """Clean and normalize extracted text.

    - Removes excessive whitespace
    - Normalizes line endings
    - Removes hidden characters
    """
    if not text:
        return ""

    # Replace multiple whitespace with single space
    import re
    text = re.sub(r"\s+", " ", text)

    # Remove non-printable characters except common ones
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Remove multiple newlines
    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


def is_readable_file(file_path: str) -> bool:
    """Check if a file can be read and parsed."""
    try:
        path = Path(file_path)
        return path.exists() and path.is_file() and path.stat().st_size > 0
    except Exception:
        return False
